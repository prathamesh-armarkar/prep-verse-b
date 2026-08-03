"""DOCX text extraction."""

from utils.pdf_parser import DocumentParsingError


def extract_docx_text(file_path):
    try:
        from docx import Document
        document = Document(file_path)
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
        cells = [cell.text.strip() for table in document.tables for row in table.rows for cell in row.cells]
        text = "\n".join(part for part in paragraphs + cells if part)
        if not text:
            raise DocumentParsingError("DOCX", "No extractable text was found in the DOCX file.")
        return text
    except DocumentParsingError:
        raise
    except Exception as exc:
        raise DocumentParsingError("DOCX", str(exc)) from exc
